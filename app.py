import streamlit as st
import tempfile
import os
import io
import time
from datetime import datetime

# Clientes de IA
from openai import OpenAI
from google import genai

# Librería para exportar a Documentos (Google Docs / Word)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Configuración de la interfaz
st.set_page_config(page_title="Panel Directivo - Relatoría CTE", page_icon="🎛️", layout="wide")
st.title("🎛️ Panel de Control Directivo: CTE USAER 2E")
st.markdown("---")

# Función para generar el documento oficial
def generar_documento_oficial(contenido_relatoria):
    doc = Document()
    
    # Fecha oficial dinámica
    meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
    hoy = datetime.now()
    fecha_str = f"Mérida, Yucatán a {hoy.day:02d} de {meses[hoy.month-1]} de {hoy.year}"
    
    # Encabezado de Fecha
    p_fecha = doc.add_paragraph()
    r_fecha = p_fecha.add_run(fecha_str)
    r_fecha.font.name = 'Arial'
    r_fecha.font.size = Pt(11)
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Título oficial
    p_titulo = doc.add_paragraph()
    r_titulo = p_titulo.add_run("RELATORÍA")
    r_titulo.bold = True
    r_titulo.font.name = 'Arial'
    r_titulo.font.size = Pt(12)
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Inserción de párrafos
    for linea in contenido_relatoria.split('\n'):
        linea_limpia = linea.strip()
        if not linea_limpia:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if linea_limpia.startswith('##') or linea_limpia.startswith('#'):
            r = p.add_run(linea_limpia.replace('#', '').strip())
            r.bold = True
            r.font.size = Pt(12)
        else:
            partes = linea_limpia.split('**')
            for i, parte in enumerate(partes):
                r = p.add_run(parte)
                if i % 2 != 0:
                    r.bold = True
        
        for run in p.runs:
            run.font.name = 'Arial'
            if not run.font.size:
                run.font.size = Pt(11)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 2. Configuración en la Barra Lateral
with st.sidebar:
    st.header("⚙️ Motor de IA y Ajustes")
    motor_ia = st.radio("Selecciona el motor de IA:", ["ChatGPT (OpenAI)", "Gemini (Google)"])
    st.markdown("---")
    num_sesion = st.selectbox("Sesión de CTE:", ["Fase Intensiva", "Primera Sesión", "Segunda Sesión", "Tercera Sesión"])
    enfoque_especial = st.text_input("Tema central:", placeholder="Ej. Ajustes razonables, BAP...")

# 3. Métodos de Captura de Audio
st.subheader("🎙️ Captura de Audio de la Plenaria")
modo_grabacion = st.radio(
    "Método de captura:", 
    ["Subir archivo de audio (Recomendado para exponer sin pausas)", "Grabar en el navegador"],
    horizontal=True
)

audio_path_temporal = None
col1, col2 = st.columns([1.5, 1])

with col1:
    if "Subir" in modo_grabacion:
        archivo_subido = st.file_uploader("Arrastra tu archivo (MP3, WAV, M4A)", type=["wav", "mp3", "m4a"])
        if archivo_subido:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
                tmp_file.write(archivo_subido.getvalue())
                audio_path_temporal = tmp_file.name
    else:
        audio_grabado = st.audio_input("Haz clic para grabar en vivo")
        if audio_grabado:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                tmp_file.write(audio_grabado.getvalue())
                audio_path_temporal = tmp_file.name

with col2:
    st.subheader("📝 Notas de Dirección")
    notas_directivo = st.text_area("Observaciones clave de la sesión (Opcional):", height=150)

# Inicializar variable de estado para el texto
if 'texto_relatoria' not in st.session_state:
    st.session_state.texto_relatoria = None

# 4. Procesamiento Inteligente
if audio_path_temporal:
    st.markdown("---")
    if st.button("🚀 Procesar Audio y Generar Relatoría Oficial", type="primary", use_container_width=True):
        
        prompt_relatoria = f"""
        Actúa como el secretario técnico de la USAER 2E. Analiza el registro de la {num_sesion} del CTE.
        Enfoque: '{enfoque_especial}'. Notas directivas: '{notas_directivo}'.

        Genera el contenido de la relatoría (sin título ni fecha, se inyectan automáticamente). Estructura el texto directamente:
        1. **Contexto y Desarrollo:** Resumen del diálogo del colegiado.
        2. **Reflexiones y Retos Pedagógicos:** Puntos críticos analizados.
        3. **Acuerdos y Compromisos:** Tareas concretas acordadas.
        Redacta en un tono directivo, formal y claro para archivo oficial.
        """

        # --- OPCIÓN 1: CHATGPT (OPENAI) ---
        if "ChatGPT" in motor_ia:
            with st.spinner("Transcribiendo con Whisper y redactando con GPT-4o..."):
                try:
                    openai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
                    with open(audio_path_temporal, "rb") as audio_file:
                        transcripcion = openai_client.audio.transcriptions.create(model="whisper-1", file=audio_file)
                    
                    response = openai_client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "Eres un redactor técnico institucional de educación especial."},
                            {"role": "user", "content": f"{prompt_relatoria}\n\nTranscripción:\n{transcripcion.text}"}
                        ]
                    )
                    st.session_state.texto_relatoria = response.choices[0].message.content
                except Exception as e:
                    st.error(f"Error con OpenAI: {e}")

        # --- OPCIÓN 2: GEMINI CON RETRY AUTOMÁTICO ANTE 503 ---
        else:
            with st.spinner("Subiendo audio y redactando con Gemini..."):
                try:
                    gemini_client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
                    audio_file = gemini_client.files.upload(file=audio_path_temporal)

                    modelos_disponibles = ['gemini-2.0-flash-lite', 'gemini-1.5-flash', 'gemini-3.6-flash']
                    response = None

                    for nombre_modelo in modelos_disponibles:
                        for intento in range(3):
                            try:
                                response = gemini_client.models.generate_content(
                                    model=nombre_modelo,
                                    contents=[prompt_relatoria, audio_file]
                                )
                                break
                            except Exception as err:
                                if "503" in str(err) or "UNAVAILABLE" in str(err):
                                    time.sleep(2 ** intento)
                                else:
                                    break
                        if response:
                            break

                    if response:
                        st.session_state.texto_relatoria = response.text
                    else:
                        st.error("Servidores de Google en alta demanda temporal. Cambia a ChatGPT en el menú lateral.")

                    gemini_client.files.delete(name=audio_file.name)
                except Exception as e:
                    st.error(f"Error con Gemini: {e}")

        # Limpieza de archivo local
        if os.path.exists(audio_path_temporal):
            os.remove(audio_path_temporal)

# 5. Visualización y Exportación
if st.session_state.texto_relatoria:
    st.markdown("---")
    st.header("📄 Vista Previa de la Relatoría")
    st.markdown(st.session_state.texto_relatoria)
    
    archivo_docx = generar_documento_oficial(st.session_state.texto_relatoria)
    
    st.markdown("### 💾 Exportar Documento")
    st.download_button(
        label="📥 Descargar formato oficial (Compatible con Google Docs / Word)",
        data=archivo_docx,
        file_name=f"Relatoria_CTE_{num_sesion.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )
